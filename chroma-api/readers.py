from typing import List, Union
import requests
import logging
from langchain.text_splitter import RecursiveCharacterTextSplitter
import docx
import PyPDF2
import io
import json
import csv
import ssl
import os
from concurrent.futures import ThreadPoolExecutor

# Отключаем проверку SSL для тестирования
ssl._create_default_https_context = ssl._create_unverified_context

class DocumentReader:
    @staticmethod
    def is_url(path: str) -> bool:
        """
        Проверяет, является ли путь URL-адресом
        """
        return path.startswith(('http://', 'https://'))

    @staticmethod
    def is_safe_path(path: str) -> bool:
        """
        Проверяет, безопасен ли путь к файлу
        """
        # Список разрешенных директорий
        allowed_dirs = [
            '/var/www/html/data',
            '/opt/data'
            # Добавьте другие разрешенные директории
        ]

        # Получаем абсолютный путь
        abs_path = os.path.abspath(path)

        # Проверяем, находится ли файл в разрешенной директории
        return any(abs_path.startswith(allowed_dir) for allowed_dir in allowed_dirs)

    @staticmethod
    def check_file_size(path: str, max_size_mb: int = 10) -> bool:
        """
        Проверяет размер файла
        """
        max_size_bytes = max_size_mb * 1024 * 1024
        return os.path.getsize(path) <= max_size_bytes

    @staticmethod
    def download_file(path: str) -> str:
        """
        Загружает файл по URL или читает локальный файл
        """
        try:
            if DocumentReader.is_url(path):
                # Если это URL, загружаем через requests
                response = requests.get(path, verify=False, timeout=30)
                response.raise_for_status()
                return response.text
            else:
                # Проверяем безопасность пути
                if not DocumentReader.is_safe_path(path):
                    raise ValueError(f"Access to path not allowed: {path}")

                # Проверяем существование файла
                if not os.path.exists(path):
                    raise FileNotFoundError(f"File not found: {path}")

                # Проверяем размер файла
                if not DocumentReader.check_file_size(path):
                    raise ValueError(f"File too large: {path}")

                # Читаем файл
                with open(path, 'r', encoding='utf-8') as file:
                    return file.read()

        except Exception as e:
            logging.error(f"Error reading file: {str(e)}")
            logging.error(f"Path: {path}")
            raise

    @staticmethod
    def download_binary(path: str) -> bytes:
        """
        Загружает бинарный файл по URL или читает локальный файл
        """
        try:
            if DocumentReader.is_url(path):
                # Если это URL, загружаем через requests
                response = requests.get(path, verify=False, timeout=30)
                response.raise_for_status()
                return response.content
            else:
                # Проверяем безопасность пути
                if not DocumentReader.is_safe_path(path):
                    raise ValueError(f"Access to path not allowed: {path}")

                # Проверяем существование файла
                if not os.path.exists(path):
                    raise FileNotFoundError(f"File not found: {path}")

                # Проверяем размер файла
                if not DocumentReader.check_file_size(path):
                    raise ValueError(f"File too large: {path}")

                # Читаем файл
                with open(path, 'rb') as file:
                    return file.read()

        except Exception as e:
            logging.error(f"Error reading binary file: {str(e)}")
            logging.error(f"Path: {path}")
            raise

    @staticmethod
    def split_text(text: str, chunk_size: int = None, chunk_overlap: int = None,
                   custom_separators: List[str] = None,
                   parallel: bool = True,
                   max_workers: int = 4) -> List[str]:
        """
        Разделяет текст на чанки параллельно
        """
        if custom_separators and any(sep in text for sep in custom_separators):
            chunks = []

            # Разбиваем текст на строки
            lines = text.split('\n')

            # Функция для обработки части строк
            def process_lines(lines_subset):
                local_chunks = []
                current = []

                for line in lines_subset:
                    is_separator = any(sep in line for sep in custom_separators)

                    if is_separator:
                        if current:
                            local_chunks.append('\n'.join(current).strip())
                            current = []
                    else:
                        current.append(line)

                if current:
                    local_chunks.append('\n'.join(current).strip())

                return local_chunks

            if parallel and len(lines) > 1000:  # Порог для параллельной обработки
                # Разбиваем строки на части для параллельной обработки
                lines_per_worker = len(lines) // max_workers
                line_chunks = [lines[i:i + lines_per_worker] for i in range(0, len(lines), lines_per_worker)]

                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    chunk_results = executor.map(process_lines, line_chunks)

                for result in chunk_results:
                    chunks.extend(result)
            else:
                chunks = process_lines(lines)

            return [chunk for chunk in chunks if chunk]

        elif chunk_size:
            chunk_overlap = chunk_overlap or 0

            # Используем RecursiveCharacterTextSplitter с оптимизированными параметрами
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", ". ", "! ", "? ", ".", "!", "?", " ", ""],
                keep_separator=False
            )

            # Если текст большой, разбиваем его на части для параллельной обработки
            if parallel and len(text) > 100000:  # Порог для параллельной обработки
                # Примерно разбиваем текст на части по параграфам
                paragraphs = text.split('\n\n')
                chunks = []

                def split_paragraph(para):
                    if para.strip():
                        return text_splitter.split_text(para)
                    return []

                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    results = executor.map(split_paragraph, paragraphs)

                for result in results:
                    chunks.extend(result)

                return chunks
            else:
                return text_splitter.split_text(text)

        return [text]

    @staticmethod
    def read_txt(path: str, **kwargs) -> str:
        """
        Читает текстовый файл
        """
        return DocumentReader.download_file(path)

    @staticmethod
    def read_docx(path: str, **kwargs) -> str:
        """
        Читает файл DOCX
        """
        try:
            content = DocumentReader.download_binary(path)
            doc = docx.Document(io.BytesIO(content))

            # Параллельное извлечение текста из параграфов
            def process_paragraphs(paragraphs):
                return [p.text for p in paragraphs]

            if kwargs.get('parallel', True) and len(doc.paragraphs) > 100:
                max_workers = kwargs.get('max_workers', 4)
                # Разбиваем параграфы на группы
                para_chunks = [doc.paragraphs[i:i + 100] for i in range(0, len(doc.paragraphs), 100)]

                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    results = executor.map(process_paragraphs, para_chunks)

                all_paragraphs = []
                for result in results:
                    all_paragraphs.extend(result)

                return '\n'.join(all_paragraphs)
            else:
                return '\n'.join(process_paragraphs(doc.paragraphs))

        except Exception as e:
            logging.error(f"Error reading DOCX file: {str(e)}")
            raise

    @staticmethod
    def read_pdf(path: str, **kwargs) -> str:
        """
        Читает файл PDF
        """
        try:
            content = DocumentReader.download_binary(path)
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            def extract_page_text(page):
                return page.extract_text()

            if kwargs.get('parallel', True) and len(pdf_reader.pages) > 10:
                max_workers = kwargs.get('max_workers', 4)

                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    text_chunks = executor.map(extract_page_text, pdf_reader.pages)

                return '\n'.join(text_chunks)
            else:
                text = []
                for page in pdf_reader.pages:
                    text.append(extract_page_text(page))
                return '\n'.join(text)

        except Exception as e:
            logging.error(f"Error reading PDF file: {str(e)}")
            raise

    @staticmethod
    def read_csv(path: str, **kwargs) -> List[str]:
        """
        Читает файл CSV и форматирует каждую строку как отдельный документ
        """
        try:
            content = DocumentReader.download_file(path)
            csv_file = io.StringIO(content)
            csv_reader = csv.DictReader(csv_file)
            rows = list(csv_reader)

            def process_row(row):
                # Форматируем каждую строку как текст с заголовками
                formatted_text = []
                for header, value in row.items():
                    formatted_text.append(f"{header}: {value}")
                return '\n'.join(formatted_text)

            if kwargs.get('parallel', True) and len(rows) > 100:
                max_workers = kwargs.get('max_workers', 4)

                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    documents = list(executor.map(process_row, rows))

                return documents
            else:
                return [process_row(row) for row in rows]

        except Exception as e:
            logging.error(f"Error reading CSV file: {str(e)}")
            raise

    @staticmethod
    def read_json(path: str, **kwargs) -> str:
        """
        Читает файл JSON
        """
        try:
            content = DocumentReader.download_file(path)
            # Пытаемся отформатировать JSON для лучшей читаемости
            json_data = json.loads(content)
            return json.dumps(json_data, ensure_ascii=False, indent=2)
        except Exception as e:
            logging.error(f"Error reading JSON file: {str(e)}")
            raise

    @staticmethod
    def read_file(path: str, **kwargs) -> Union[str, List[str]]:
        """
        Читает файл любого поддерживаемого формата и разделяет на чанки
        """
        try:
            # Определяем тип файла по расширению
            file_type = path.split('.')[-1].lower()

            # Словарь поддерживаемых форматов и соответствующих методов
            readers = {
                'txt': DocumentReader.read_txt,
                'docx': DocumentReader.read_docx,
                'pdf': DocumentReader.read_pdf,
                'csv': DocumentReader.read_csv,
                'json': DocumentReader.read_json
            }

            # Проверяем поддержку формата
            if file_type not in readers:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Читаем содержимое файла
            content = readers[file_type](path, **kwargs)

            # Если это список (например, для CSV), возвращаем как есть
            if isinstance(content, list):
                return content

            # Разделяем на чанки, если нужно
            return DocumentReader.split_text(
                content,
                chunk_size=kwargs.get('chunk_size'),
                chunk_overlap=kwargs.get('chunk_overlap'),
                custom_separators=kwargs.get('custom_separators'),
                parallel=kwargs.get('parallel', True),
                max_workers=kwargs.get('max_workers', 4)
            )

        except Exception as e:
            logging.error(f"Error in read_file: {str(e)}")
            logging.error(f"Path: {path}")
            logging.error(f"kwargs: {kwargs}")
            raise